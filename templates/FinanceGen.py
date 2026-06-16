# -*- coding: utf-8 -*-
"""生成财务预测表docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import io

def build():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    def hd(t, lv=1):
        h = doc.add_heading(t, level=lv)
        for r in h.runs: r.font.color.rgb = RGBColor(0x8B,0x00,0x00)
    def p(txt):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(txt); r.font.size = Pt(12)
        r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    def tb(hds, rows):
        t = doc.add_table(rows=1+len(rows), cols=len(hds))
        t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(hds):
            c = t.rows[0].cells[j]; c.text = h
            for par in c.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in par.runs: r.font.bold = True; r.font.size = Pt(11)
        for i, row in enumerate(rows):
            for j, v in enumerate(row):
                c = t.rows[i+1].cells[j]; c.text = str(v)
                for par in c.paragraphs:
                    for r in par.runs: r.font.size = Pt(11)

    hd('未来五年营收预测表（万元）', 1)
    tb(['项目', '2026年', '2027年', '2028年', '2029年', '2030年'], [
        ['高校订阅收入', '3.84', '19.2', '51.2', '80.0', '120.0'],
        ['企业服务收入', '0', '2.0', '8.0', '20.0', '50.0'],
        ['散客订阅收入', '0', '0.5', '2.0', '6.0', '15.0'],
        ['学校基金或课题', '0.5', '1.0', '1.0', '1.0', '1.0'],
        ['营业收入合计', '4.34', '22.7', '62.2', '107.0', '186.0'],
        ['减：营业成本', '0.5', '2.0', '5.0', '12.0', '25.0'],
        ['销售推广费用', '1.5', '5.0', '12.0', '22.0', '38.0'],
        ['管理及其他费用', '1.0', '3.0', '8.0', '15.0', '28.0'],
        ['净利润', '1.34', '12.7', '37.2', '58.0', '95.0'],
    ])
    doc.add_paragraph()
    hd('费用预测表（万元）', 2)
    tb(['费用项目', '2026年', '2027年', '2028年', '2029年', '2030年'], [
        ['服务器及带宽', '0.3', '1.0', '2.5', '6.0', '12.0'],
        ['数据库维护及扩充', '0.2', '0.5', '1.5', '3.0', '6.0'],
        ['论文发表及学术推广', '0.5', '1.5', '3.0', '5.0', '8.0'],
        ['线上推广及广告', '0.3', '1.0', '2.5', '5.0', '10.0'],
        ['差旅及会议', '0.3', '1.0', '2.5', '4.0', '6.0'],
        ['人员劳务', '0.5', '3.0', '8.0', '18.0', '35.0'],
        ['其他杂项', '0.2', '0.5', '1.0', '3.0', '6.0'],
        ['合计', '2.3', '8.5', '21.0', '44.0', '83.0'],
    ])
    doc.add_paragraph()
    hd('营收测算依据', 2)
    p('高校订阅收入方面，以每校每年12800元的定价为基准。2026年以贵州省内中医药院校为起点，计划覆盖贵州中医药大学及省内2至3所院校，按3所签约计算收入3.84万元。2027年扩展至全国15所院校，2028年扩展至40所，2029年80所，2030年覆盖全国主要中医药院校约120所，按百分之八十续约率估算为120所以上。高校客户是平台最稳定的收入来源，合同周期通常为一年，续约率预期较高。')
    p('企业服务收入方面，2026年为平台建设期，暂不拓展企业客户。2027年起与中药企业建立合作，从中小型企业的每年2980至6980元起步。2028年企业客户数量增长至30至40家，2029年增长至60至80家，2030年覆盖约150家规模以上中药企业。企业端收入增长快于高校端，但占比相对较小。')
    p('散客订阅收入方面，基础版免费使用培养用户习惯后，部分活跃用户转化为付费订阅。2027年起按周卡18元、月卡30元、年卡280元的价格体系，以活跃用户百分之三至五的转化率估算。2028年后随着用户基数增长，散客收入逐步成为补充收入来源。')
    p('营业外收入主要为学校创新创业基金、课题经费支持和各级大赛奖金，每年约0.5至1万元。')
    p('成本端，平台全部技术栈为开源软件，无商业授权依赖，营业成本主要为服务器租赁和带宽费用，随用户规模线性增长。销售推广费用包含论文发表版面费、学术会议注册费和差旅费、线上广告投放等。管理费用包括人员劳务、行政办公等。由于开源技术栈的零边际成本特性，综合利润率从第一年的百分之三十一提升至第五年的百分之五十一，成熟期稳定在百分之五十以上。')
    p('项目初期投入1万元即可启动（自有资金0.5万元加学校创新基金0.5万元），第一年即实现正向净利润1.34万元，投资回收期低于一年。')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
